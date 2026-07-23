import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class TextExtractor:
    """Native text extraction from text-layer PDFs, DOCX, XLSX, TXT files.

    Always try this FIRST before OCR — text extraction is instant (<1s)
    vs OCR (5-30s). Only fall back to OCR if extracted text is empty.
    """

    def extract(self, file_path: str) -> dict:
        path = Path(file_path)
        if not path.exists():
            return {"full_text": "", "text_lines": [], "tables": [], "error": "file_not_found"}

        ext = path.suffix.lower()
        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".xlsx": self._extract_xlsx,
            ".txt": self._extract_txt,
        }

        extractor = extractors.get(ext)
        if not extractor:
            return {"full_text": "", "text_lines": [], "tables": [], "error": "unsupported_format"}

        try:
            return extractor(path)
        except Exception as e:
            logger.warning("Native extraction failed for %s: %s", file_path, e)
            return {"full_text": "", "text_lines": [], "tables": [], "error": str(e)}

    def _extract_pdf(self, path: Path) -> dict:
        try:
            import pdfplumber
        except ImportError:
            return {"full_text": "", "text_lines": [], "tables": [], "error": "pdfplumber not installed"}

        with pdfplumber.open(str(path)) as pdf:
            text_lines = []
            all_tables = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_lines.extend(text.splitlines())
                tables = page.extract_tables()
                if tables:
                    for t in tables:
                        if t:
                            all_tables.append([row for row in t if any(cell is not None and cell.strip() for cell in row)])

            full_text = "\n".join(text_lines)
            return {
                "full_text": full_text,
                "text_lines": text_lines,
                "tables": all_tables,
                "source_type": "TEXT_PDF",
            }

    def _extract_docx(self, path: Path) -> dict:
        try:
            from docx import Document
        except ImportError:
            return {"full_text": "", "text_lines": [], "tables": [], "error": "python-docx not installed"}

        doc = Document(str(path))
        text_lines = [p.text for p in doc.paragraphs if p.text.strip()]
        all_tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(c for c in cells):
                    rows.append(cells)
            if rows:
                all_tables.append(rows)

        full_text = "\n".join(text_lines)
        return {
            "full_text": full_text,
            "text_lines": text_lines,
            "tables": all_tables,
            "source_type": "NATIVE_DOCX",
        }

    def _extract_xlsx(self, path: Path) -> dict:
        try:
            import openpyxl
        except ImportError:
            return {"full_text": "", "text_lines": [], "tables": [], "error": "openpyxl not installed"}

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        text_lines = []
        all_tables = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                row_text = " ".join(c for c in cells if c)
                if row_text.strip():
                    text_lines.append(row_text)
                if any(c for c in cells):
                    rows.append(cells)
            if rows:
                all_tables.append(rows)
        wb.close()

        full_text = "\n".join(text_lines)
        return {
            "full_text": full_text,
            "text_lines": text_lines,
            "tables": all_tables,
            "source_type": "NATIVE_XLSX",
        }

    def _extract_txt(self, path: Path) -> dict:
        text = path.read_text(encoding="utf-8", errors="replace")
        text_lines = [l for l in text.splitlines() if l.strip()]
        return {
            "full_text": text,
            "text_lines": text_lines,
            "tables": [],
            "source_type": "NATIVE_TXT",
        }
